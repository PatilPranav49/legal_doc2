"""
layer1_extraction/docx_extractor.py
Extract text from .docx files using python-docx.
"""
from pathlib import Path
from utils.logger import get_logger

logger = get_logger(__name__)


def extract_docx_text(file_path: str) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {
            "text": "",
            "pages": 0,
            "success": False,
            "error": "python-docx not installed. Run: pip install python-docx"
        }

    try:
        doc = Document(file_path)

        #  Faster text extraction
        paragraphs = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                paragraphs.append(txt)

        text_parts = paragraphs

        #   Efficient table extraction
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))

        text = "\n\n".join(text_parts)

        #   Avoid heavy split on large text
        word_count = text.count(" ") + 1 if text else 0
        estimated_pages = max(1, word_count // 300)

        logger.info(f"DOCX extracted: {len(text)} chars, ~{estimated_pages} pages")

        return {
            "text": text,
            "pages": estimated_pages,
            "success": True,
            "error": None,
        }

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return {
            "text": "",
            "pages": 0,
            "success": False,
            "error": str(e)
        }